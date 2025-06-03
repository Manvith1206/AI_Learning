from typing import Dict, List, Any, Tuple, Optional
from app.domain.models import Document, DocumentChunk
from app.domain.services import RAGService
from app.core.events.event_bus import event_bus
from app.core.events.event_handlers import DocumentProcessedEvent
import uuid
import os


class DocumentProcessingUseCase:
    """Application use case for document processing"""
    
    def __init__(self, rag_service: RAGService):
        self.rag_service = rag_service
    
    def extract_text_from_file(self, file_obj) -> str:
        """Extract text from an uploaded file"""
        # Get file extension
        file_name = file_obj.name
        file_extension = os.path.splitext(file_name)[1].lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_obj)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_obj)
        elif file_extension == '.txt':
            return self._extract_from_txt(file_obj)
        elif file_extension == '.csv':
            return self._extract_from_csv(file_obj)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_obj) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(file_obj)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_obj) -> str:
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(file_obj)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_obj) -> str:
        """Extract text from TXT file"""
        try:
            text = file_obj.read().decode('utf-8')
            return text
        except Exception as e:
            raise ValueError(f"Error extracting text from TXT: {str(e)}")
    
    def _extract_from_csv(self, file_obj) -> str:
        """Extract text from CSV file"""
        try:
            import pandas as pd
            df = pd.read_csv(file_obj)
            return df.to_string()
        except Exception as e:
            raise ValueError(f"Error extracting text from CSV: {str(e)}")
    
    def process_document(self, file_obj) -> Tuple[Document, List[DocumentChunk]]:
        """Process a document file and return the document and its chunks"""
        # Extract text from file
        text = self.extract_text_from_file(file_obj)
        
        # Create document
        document = Document(
            id=str(uuid.uuid4()),
            content=text,
            metadata={
                "filename": file_obj.name,
                "file_size": file_obj.size,
                "file_type": os.path.splitext(file_obj.name)[1].lower()
            }
        )
        
        # Process document
        processed_document = self.rag_service.process_document(document)
        
        # Publish event
        event_bus.publish(DocumentProcessedEvent(
            document_id=document.id,
            num_chunks=len(document.chunks),
            metadata=document.metadata
        ))
        
        return processed_document, processed_document.chunks
